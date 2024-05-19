from docx import Document
from docx.shared import Inches
import os

def insert_images_to_word(doc_path, images_dir, image_prefix, start_image, end_image):
    # Create a new Word document
    doc = Document()

    # Loop through the specified range of image numbers
    for i in range(start_image, end_image + 1):
        image_path = os.path.join(images_dir, f"{image_prefix}_{i}.jpg")
        if os.path.exists(image_path):
            doc.add_picture(image_path, width=Inches(4.0))  # Insert image
            doc.add_paragraph()  # Add a space between images
        else:
            print(f"Image not found: {image_path}")

    # Save the document
    doc.save(doc_path)
    print(f"Document saved: {doc_path}")

if __name__ == '__main__':
    doc_dir = 'D:/Cache/ztzp-ConCaSys/TestFuction/test/pic2'
    images_dir = 'D:/Cache/ztzp-ConCaSys/TestFuction/test/pic2'
    image_prefix = '640'

    # Insert images for the '问题.docx' document
    problem_doc_path = os.path.join(doc_dir, '问题-答案.docx')
    insert_images_to_word(problem_doc_path, images_dir, image_prefix, 2, 10)

